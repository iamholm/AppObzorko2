# Файл: docx_to_excel_processor.py
# Содержит класс DocxToExcelProcessor для обработки DOCX файлов и конвертации в Excel

from docx import Document
import openpyxl
from openpyxl.utils import get_column_letter
import re
from datetime import datetime

class DocxToExcelProcessor:
    """
    Класс для обработки документов DOCX и преобразования их в Excel
    с дополнительной обработкой данных.
    """
    
    def convert_docx_to_excel(self, docx_path, excel_path):
        """Извлечение таблиц из DOCX и сохранение в Excel"""
        # Открываем DOCX-файл
        document = Document(docx_path)
        
        # Проверяем, есть ли таблицы
        if not document.tables:
            return 0
        
        # Создаем новую рабочую книгу Excel
        workbook = openpyxl.Workbook()
        # Удаляем стандартный лист
        default_sheet = workbook.active
        workbook.remove(default_sheet)
        
        # Счетчик таблиц
        table_count = 0
        
        # Для каждой таблицы из docx
        for i, table in enumerate(document.tables):
            # Создаем новый лист для каждой таблицы
            sheet_name = f"Таблица_{i+1}"
            sheet = workbook.create_sheet(title=sheet_name)
            table_count += 1
            
            # Копируем данные из таблицы docx в Excel
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    # Excel использует индексацию с 1, а не с 0
                    excel_row = row_idx + 1
                    excel_col = col_idx + 1
                    sheet.cell(row=excel_row, column=excel_col).value = cell.text
            
            # Автоподбор ширины столбцов
            self._adjust_column_width(sheet)
        
        # Сохраняем Excel-файл
        workbook.save(excel_path)
        
        return table_count
    
    def process_excel_file(self, excel_path):
        """Удаление столбцов A и C из Excel-файла и обработка первой строки"""
        # Загружаем рабочую книгу
        workbook = openpyxl.load_workbook(excel_path)
        
        # Колонки для удаления в обратном порядке (C, A)
        # Важно: удаляем сначала большие индексы, потом меньшие,
        # чтобы не смещались индексы колонок при удалении
        columns_to_remove = [3, 1]  # C = 3, A = 1
        
        stats = {
            "sheets_processed": 0,
            "rows_deleted": 0,
            "dates_normalized": 0,
            "birth_dates_normalized": 0,
            "end_dates_normalized": 0,
            "text_moved": 0,
            "court_info_moved": 0,
            "court_dates_normalized": 0,
            "formatted_cells": 0
        }
        
        # Обрабатываем каждый лист
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            stats["sheets_processed"] += 1
            
            # ВАЖНО: Сначала проверяем, нужно ли удалить первую строку
            # Получаем значение ВТОРОЙ ячейки (B1) для проверки
            second_cell_value = sheet.cell(row=1, column=2).value
            
            # Определяем, нужно ли удалять первую строку
            delete_first_row = not self._is_date(second_cell_value)
            
            # Удаляем столбцы
            for col_idx in columns_to_remove:
                sheet.delete_cols(col_idx, 1)
            
            # Теперь удаляем первую строку, если нужно
            if delete_first_row:
                sheet.delete_rows(1, 1)
                stats["rows_deleted"] += 1
            
            # Нормализуем даты в первом столбце (бывший B, теперь A после удаления)
            normalized_count = self._normalize_dates(sheet, 1)  # Столбец 1 (A)
            stats["dates_normalized"] += normalized_count
            
            # Нормализуем даты рождения в третьем столбце (бывший E, теперь C после удаления столбцов A и C)
            birth_normalized_count = self._normalize_birth_dates(sheet, 3)  # Столбец 3 (C)
            stats["birth_dates_normalized"] += birth_normalized_count
            
            # Обрабатываем столбец 8 (бывший J, теперь H/6 после удаления столбцов A и C)
            end_dates_count, moved_text_count = self._process_end_dates(sheet, 6, 8)  # Столбец 6 (F) и 8 (H)
            stats["end_dates_normalized"] += end_dates_count
            stats["text_moved"] += moved_text_count
            
            # Обрабатываем столбцы 4 и 5 (бывшие F и G, новые D и E) и ищем информацию о судах
            court_moved = self._move_court_info(sheet, source_columns=(4, 5), target_column=9)
            stats["court_info_moved"] += court_moved
            
            # Нормализуем даты в столбце с информацией о судах
            court_normalized = self._normalize_dates_in_court_info(sheet, 9)
            stats["court_dates_normalized"] += court_normalized
            
            # Форматируем информацию о судах для улучшения читаемости
            formatted_cells = self._format_court_info(sheet, 9)
            stats["formatted_cells"] += formatted_cells
            
            # Автоподбор ширины столбцов
            self._adjust_column_width(sheet)
        
        # Сохраняем изменения
        workbook.save(excel_path)
        
        # Общее количество нормализованных дат
        stats["total_dates_normalized"] = stats["dates_normalized"] + stats["birth_dates_normalized"] + stats["end_dates_normalized"] + stats["court_dates_normalized"]
        
        return stats
    
    def _format_court_info(self, sheet, column_index=9):
        """
        Улучшает читаемость информации о судах в указанном столбце
        
        Возвращает количество отформатированных ячеек
        """
        formatted_count = 0
        
        # Обрабатываем все ячейки в указанном столбце
        for row in range(1, sheet.max_row + 1):
            cell = sheet.cell(row=row, column=column_index)
            value = cell.value
            
            # Пропускаем пустые ячейки
            if not value:
                continue
                
            value_str = str(value).strip()
            
            # Применяем форматирование
            formatted_text = self._apply_formatting_rules(value_str)
            
            # Если текст изменился, обновляем ячейку
            if formatted_text != value_str:
                cell.value = formatted_text
                formatted_count += 1
        
        return formatted_count
    
    def _apply_formatting_rules(self, text):
        """
        Применяет базовое форматирование к тексту
        
        Возвращает отформатированный текст
        """
        # Убираем множественные пробелы
        text = re.sub(r'\s+', ' ', text)
        
        # Добавляем пробелы после знаков препинания
        text = re.sub(r'([.,;:])(?!\s)', r'\1 ', text)
        
        # Убираем пробелы перед знаками препинания
        text = re.sub(r'\s+([.,;:])', r'\1', text)
        
        # Добавляем точку в конце предложения, если ее нет
        if text and not text.endswith(('.', '!', '?')):
            text += '.'
        
        # Первая буква предложения должна быть заглавной
        if text:
            text = text[0].upper() + text[1:]
        
        return text
    
    def _normalize_dates_in_court_info(self, sheet, column_index=9):
        """
        Нормализует все даты в столбце с информацией о судах к формату ДД.ММ.ГГГГ
        
        Возвращает количество нормализованных дат
        """
        normalized_count = 0
        
        # Обрабатываем все ячейки в указанном столбце
        for row in range(1, sheet.max_row + 1):
            cell = sheet.cell(row=row, column=column_index)
            value = cell.value
            
            # Пропускаем пустые ячейки
            if not value:
                continue
                
            value_str = str(value).strip()
            
            # Извлекаем все даты из текста
            dates = self._extract_all_dates_from_text(value_str)
            
            # Если даты найдены
            if dates:
                modified_value = value_str
                # Для каждой найденной даты
                for date in dates:
                    # Нормализуем дату
                    normalized_date = self._parse_and_normalize_date(date)
                    if normalized_date:
                        # Заменяем исходную дату на нормализованную
                        modified_value = modified_value.replace(date, normalized_date)
                        normalized_count += 1
                
                # Обновляем значение ячейки только если были изменения
                if modified_value != value_str:
                    cell.value = modified_value
        
        return normalized_count
    
    def _move_court_info(self, sheet, source_columns=(4, 5), target_column=9):
        """
        Проверяет столбцы source_columns на наличие информации о судах
        и перемещает эту информацию в target_column
        
        Возвращает количество перемещенных записей о судах
        """
        moved_count = 0
        
        # Ключевые слова и паттерны для определения информации о судах
        court_keywords = [
            r'\d{2}\.\d{2}\.\d{4}.*?суд',        # Дата + суд
            r'суд.*?по ст',                       # суд + статья
            r'р/с',                               # районный суд (сокращение)
            r'г/с',                               # городской суд (сокращение)
            r'судом',                             # слово "судом"
            r'осужденный',                        # слово "осужденный"
            r'постановлением',                    # слово "постановлением"
            r'УК РФ',                             # отсылка к УК РФ
            r'л/св',                              # лишение свободы
            r'ст\. \d{1,3}',                      # статья (например, "ст. 158")
            r'ИС \d+ (год|г|лет)',                # испытательный срок
            r'Мировым судьей',                    # Мировой судья
            r'МССУ',                              # МССУ (мировой судебный участок)
        ]
        
        # Обрабатываем все строки в указанных столбцах
        for row in range(1, sheet.max_row + 1):
            for col_idx in source_columns:
                cell = sheet.cell(row=row, column=col_idx)
                value = cell.value
                
                # Пропускаем пустые ячейки
                if not value:
                    continue
                    
                value_str = str(value).strip()
                
                # Проверяем, содержит ли ячейка информацию о суде
                is_court_info = False
                
                # Проверяем наличие ключевых слов/шаблонов
                for pattern in court_keywords:
                    if re.search(pattern, value_str, re.IGNORECASE):
                        is_court_info = True
                        break
                
                # Если нашли информацию о суде
                if is_court_info:
                    # Перемещаем информацию в целевой столбец
                    target_cell = sheet.cell(row=row, column=target_column)
                    
                    # Если в целевой ячейке уже есть информация, добавляем через пробел
                    if target_cell.value:
                        target_cell.value = f"{target_cell.value} {value_str}"
                    else:
                        target_cell.value = value_str
                    
                    # Очищаем исходную ячейку
                    cell.value = ""
                    
                    moved_count += 1
        
        return moved_count
    
    def _normalize_dates(self, sheet, column_index=1):
        """
        Нормализует даты в указанном столбце к формату ДД.ММ.ГГГГ
        
        Возвращает количество нормализованных дат
        """
        normalized_count = 0
        
        # Обрабатываем все ячейки в указанном столбце
        for row in range(1, sheet.max_row + 1):
            cell = sheet.cell(row=row, column=column_index)
            value = cell.value
            
            # Пропускаем пустые ячейки
            if not value:
                continue
                
            value_str = str(value).strip()
            
            # Проверяем различные форматы дат и преобразуем их
            normalized_date = self._parse_and_normalize_date(value_str)
            
            if normalized_date:
                cell.value = normalized_date
                normalized_count += 1
        
        return normalized_count
    
    def _normalize_birth_dates(self, sheet, column_index=3):
        """
        Нормализует даты рождения в указанном столбце к формату ДД.ММ.ГГГГ
        Учитывает дополнительный текст, типа "г.р."
        
        Возвращает количество нормализованных дат
        """
        normalized_count = 0
        
        # Обрабатываем все ячейки в указанном столбце
        for row in range(1, sheet.max_row + 1):
            cell = sheet.cell(row=row, column=column_index)
            value = cell.value
            
            # Пропускаем пустые ячейки
            if not value:
                continue
                
            value_str = str(value).strip()
            
            # Извлекаем дату из текста, удаляя посторонние символы и тексты
            date_only = self._extract_date_from_text(value_str)
            
            if date_only:
                # Нормализуем извлеченную дату
                normalized_date = self._parse_and_normalize_date(date_only)
                if normalized_date:
                    cell.value = normalized_date
                    normalized_count += 1
        
        return normalized_count
    
    def _process_end_dates(self, sheet, date_column_index=6, text_column_index=8):
        """
        Обрабатывает столбец с датами окончания срока.
        - Если есть две даты, оставляет только вторую
        - Если есть дата и текст, оставляет дату, а текст перемещает в указанный столбец
        - Если есть только текст, перемещает его в указанный столбец и очищает ячейку
        
        Возвращает (количество нормализованных дат, количество перемещенных текстовых блоков)
        """
        normalized_count = 0
        moved_text_count = 0
        
        # Обрабатываем все ячейки в указанном столбце
        for row in range(1, sheet.max_row + 1):
            cell = sheet.cell(row=row, column=date_column_index)
            value = cell.value
            
            # Пропускаем пустые ячейки
            if not value:
                continue
                
            value_str = str(value).strip()
            
            # Проверяем наличие двух дат или даты с текстом
            dates = self._extract_all_dates_from_text(value_str)
            
            if len(dates) == 2:
                # Если нашли две даты, оставляем только вторую
                second_date = dates[1]
                normalized_date = self._parse_and_normalize_date(second_date)
                if normalized_date:
                    cell.value = normalized_date
                    normalized_count += 1
            elif len(dates) == 1:
                # Нашли одну дату, нормализуем ее
                date = dates[0]
                normalized_date = self._parse_and_normalize_date(date)
                
                # Извлекаем текст, который следует за датой
                text_after_date = value_str[value_str.find(date) + len(date):].strip()
                
                if text_after_date:
                    # Перемещаем текст в указанный столбец
                    text_cell = sheet.cell(row=row, column=text_column_index)
                    text_cell.value = text_after_date
                    moved_text_count += 1
                
                if normalized_date:
                    cell.value = normalized_date
                    normalized_count += 1
            else:
                # Если даты не нашли, но есть текст - перемещаем его
                if value_str:
                    text_cell = sheet.cell(row=row, column=text_column_index)
                    text_cell.value = value_str
                    cell.value = ""  # Очищаем исходную ячейку
                    moved_text_count += 1
        
        return normalized_count, moved_text_count
    
    def _extract_all_dates_from_text(self, text):
        """
        Извлекает все даты из текста в виде списка
        
        Например, из "14.07.25 14.08.25" извлечет ["14.07.25", "14.08.25"]
        Также обрабатывает даты с пробелами: "13. 05. 2024" -> ["13. 05. 2024"]
        """
        # Шаблоны для поиска дат в тексте
        date_patterns = [
            r'\d{1,2}\s*\.\s*\d{1,2}\s*\.\s*\d{2,4}',  # ДД.ММ.ГГ или ДД.ММ.ГГГГ (с возможными пробелами)
            r'\d{1,2}\s*/\s*\d{1,2}\s*/\s*\d{2,4}',    # ДД/ММ/ГГ или ДД/ММ/ГГГГ (с возможными пробелами)
            r'\d{1,2}\s*-\s*\d{1,2}\s*-\s*\d{2,4}',    # ДД-ММ-ГГ или ДД-ММ-ГГГГ (с возможными пробелами)
            r'\d{2}\d{2}\s*\.\s*\d{2}',                # ДДММ.ГГ (с возможными пробелами)
            r'\d{8}'                                   # ДДММГГГГ
        ]
        
        # Список для найденных дат
        found_dates = []
        
        # Ищем все даты в тексте по шаблонам
        for pattern in date_patterns:
            matches = re.findall(pattern, text)
            found_dates.extend(matches)
        
        # Удаляем возможные дубликаты
        return list(dict.fromkeys(found_dates))
        
    def _extract_date_from_text(self, text):
        """
        Извлекает первую дату из текста
        
        Возвращает только дату, если она найдена, иначе None
        """
        # Шаблоны для поиска даты в тексте
        date_patterns = [
            r'(\d{1,2}\s*\.\s*\d{1,2}\s*\.\s*\d{2,4})',  # ДД.ММ.ГГ или ДД.ММ.ГГГГ с пробелами
            r'(\d{1,2}/\d{1,2}/\d{2,4})',    # ДД/ММ/ГГ или ДД/ММ/ГГГГ
            r'(\d{1,2}-\d{1,2}-\d{2,4})',    # ДД-ММ-ГГ или ДД-ММ-ГГГГ
            r'(\d{2}\d{2}\.\d{2})',          # ДДММ.ГГ
            r'(\d{8})'                       # ДДММГГГГ
        ]
        
        # Ищем первую дату в тексте по шаблонам
        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1)
        
        return None
    
    def _parse_and_normalize_date(self, date_str):
        """
        Парсит различные форматы дат и нормализует их к формату ДД.ММ.ГГГГ
        
        Поддерживает:
        - ДД.ММ.ГГ -> ДД.ММ.ГГГГ
        - ДД.ММ.ГГГГ (оставляем как есть)
        - ДД. ММ. ГГ -> ДД.ММ.ГГГГ (с пробелами)
        - ДД. ММ. ГГГГ -> ДД.ММ.ГГГГ (с пробелами)
        - ДДММ.ГГ -> ДД.ММ.ГГГГ (пропущена точка)
        - ДД/ММ/ГГ -> ДД.ММ.ГГГГ
        - ДД-ММ-ГГ -> ДД.ММ.ГГГГ
        - ДДММГГГГ -> ДД.ММ.ГГГГ (без разделителей)
        """
        # Удаляем пробелы перед обработкой, чтобы упростить регулярные выражения
        clean_date_str = re.sub(r'\s+', '', date_str)
        
        # Проверяем различные форматы даты
        
        # Формат ДД.ММ.ГГ (двузначный год)
        match = re.match(r'^(\d{1,2})\.(\d{1,2})\.(\d{2})$', clean_date_str)
        if match:
            day, month, year = match.groups()
            full_year = self._expand_year(year)
            return f"{int(day):02d}.{int(month):02d}.{full_year}"
        
        # Формат ДД.ММ.ГГГГ (четырехзначный год)
        match = re.match(r'^(\d{1,2})\.(\d{1,2})\.(\d{4})$', clean_date_str)
        if match:
            day, month, year = match.groups()
            return f"{int(day):02d}.{int(month):02d}.{year}"
        
        # Формат ДДММ.ГГ (пропущена точка между днем и месяцем)
        match = re.match(r'^(\d{2})(\d{2})\.(\d{2})$', clean_date_str)
        if match:
            day, month, year = match.groups()
            full_year = self._expand_year(year)
            return f"{int(day):02d}.{int(month):02d}.{full_year}"
        
        # Формат ДДММГГГГ (без разделителей)
        match = re.match(r'^(\d{2})(\d{2})(\d{4})$', clean_date_str)
        if match:
            day, month, year = match.groups()
            return f"{int(day):02d}.{int(month):02d}.{year}"
        
        # Формат ДД/ММ/ГГ
        match = re.match(r'^(\d{1,2})/(\d{1,2})/(\d{2})$', clean_date_str)
        if match:
            day, month, year = match.groups()
            full_year = self._expand_year(year)
            return f"{int(day):02d}.{int(month):02d}.{full_year}"
        
        # Формат ДД/ММ/ГГГГ
        match = re.match(r'^(\d{1,2})/(\d{1,2})/(\d{4})$', clean_date_str)
        if match:
            day, month, year = match.groups()
            return f"{int(day):02d}.{int(month):02d}.{year}"
        
        # Формат ДД-ММ-ГГ
        match = re.match(r'^(\d{1,2})-(\d{1,2})-(\d{2})$', clean_date_str)
        if match:
            day, month, year = match.groups()
            full_year = self._expand_year(year)
            return f"{int(day):02d}.{int(month):02d}.{full_year}"
        
        # Формат ДД-ММ-ГГГГ
        match = re.match(r'^(\d{1,2})-(\d{1,2})-(\d{4})$', clean_date_str)
        if match:
            day, month, year = match.groups()
            return f"{int(day):02d}.{int(month):02d}.{year}"
        
        # Если ни один из форматов не подошел, возвращаем None
        return None
    
    def _expand_year(self, year_str):
        """
        Преобразует двузначный год в четырехзначный
        Правило: 00-25 -> 2000-2025, 26-99 -> 1926-1999
        """
        year = int(year_str)
        
        # Определяем текущий год для расчета порога преобразования
        current_year = datetime.now().year
        current_short_year = current_year % 100
        
        if year <= current_short_year:
            return 2000 + year
        else:
            return 1900 + year
    
    def _is_date(self, value):
        """
        Проверяет, является ли значение ячейки датой
        """
        if not value:
            return False
            
        # Преобразуем значение в строку и удаляем пробелы
        value_str = str(value).strip()
        
        # Шаблоны для проверки дат в различных форматах, включая форматы с пробелами
        date_patterns = [
            r'^\d{1,2}\s*\.\s*\d{1,2}\s*\.\s*\d{2,4}',  # ДД.ММ.ГГ или ДД.ММ.ГГГГ с пробелами
            r'^\d{1,2}/\d{1,2}/\d{2,4}',    # ДД/ММ/ГГ или ДД/ММ/ГГГГ
            r'^\d{1,2}-\d{1,2}-\d{2,4}',    # ДД-ММ-ГГ или ДД-ММ-ГГГГ
            r'^\d{2}\d{2}\.\d{2}',          # ДДММ.ГГ (пропущена точка между днем и месяцем)
            r'^\d{8}'                       # ДДММГГГГ (без разделителей)
        ]
        
        for pattern in date_patterns:
            if re.match(pattern, value_str):
                return True
                
        return False
    
    def _adjust_column_width(self, sheet):
        """Автоподбор ширины столбцов"""
        for column_cells in sheet.columns:
            max_length = 0
            column = column_cells[0].column_letter
            for cell in column_cells:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            sheet.column_dimensions[column].width = adjusted_width
